from bs4 import BeautifulSoup
import requests
import docx

#GLOBAL
BAN = ["In-IDE", "SaaS", "Self-Hosted"]
RULES = []

SONAR_URL = input("[ > ] SonarSource URL (ex: https://rules.sonarsource.com/csharp): ")
DOCUMENT_NAME = input("[ > ] Document Name (ex: JavaScript): ")
if not SONAR_URL.startswith("https://rules.sonarsource.com/"):
    print("[ - ] Please enter a valid SonarSource URL.")
    exit()

try:
    soup = BeautifulSoup(requests.get(SONAR_URL).text, "html.parser")
    for i in soup.find_all("h3"):
        if i.text not in BAN and "Unique rules to find Bugs" not in i.text:
            RULES.append(f'{i.text}:{i.find_next_sibling("span").text}')
except Exception as e:
    print(f"[ - ] Error: {e}")
    exit()

doc = docx.Document()
doc.add_heading(DOCUMENT_NAME, 0)
table = doc.add_table(rows=1, cols=3)
row = table.rows[0].cells
row[0].text = "Titre"
row[1].text = "Tag"
row[2].text = "Outil"

for i in RULES:
    row = table.add_row().cells
    row[0].text = i.split(":")[0]
    row[1].text = i.split(":")[1]
    row[2].text = "Sonar"

table.style = "Table Grid"
doc.save(f"{DOCUMENT_NAME}.docx")